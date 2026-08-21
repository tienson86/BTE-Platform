"""G3-01 clean-runtime smoke: health, version, Analyze, PDF/DOCX, ten-case probe.

Does not write engines or change Gate-2 semantics. Requires repo root on PYTHONPATH.
"""

from __future__ import annotations

import json
import sys
import time
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(ROOT))
sys.path.insert(0, str(ROOT / "release" / "gate_02"))

from fastapi.testclient import TestClient  # noqa: E402
from applications.api.app import app  # noqa: E402
from applications.api.services.customer_export import (  # noqa: E402
    export_customer_file,
    prepare_customer_report_input,
)
from _g2_01_binding_probe import CASES, COMPARE_KEYS, FROZEN, fingerprint  # noqa: E402

OUT = Path(__file__).resolve().parent / "G3_01_SMOKE.json"
DUNG = "Ngô Đắc Dũng"
HY = "Chưa đủ căn cứ xác định Hỷ thần bổ trợ riêng"


def main() -> int:
    frozen_by = {row["case"]: row for row in json.loads(FROZEN.read_text(encoding="utf-8"))}
    client = TestClient(app)
    health = client.get("/health").json()
    version = client.get("/version").json()
    api_health = client.get("/api/v1/health").json()

    mismatches: list[dict[str, object]] = []
    rows: list[dict[str, object]] = []
    dung_payload: dict | None = None
    dung_spec: dict | None = None
    for index, spec in enumerate(CASES):
        name = str(spec["name"])
        body = {k: v for k, v in spec.items() if k != "name"}
        body["full_name"] = name
        request_id = f"g3-01-{index}"
        started = time.perf_counter()
        response = client.post("/api/v1/analyze", json=body, headers={"X-Request-ID": request_id})
        elapsed = round(time.perf_counter() - started, 3)
        response.raise_for_status()
        envelope = response.json()
        data = envelope["data"]
        live = fingerprint(name, data)
        frozen = frozen_by[name]
        diffs = {
            key: {"frozen": frozen.get(key), "live": live.get(key)}
            for key in COMPARE_KEYS
            if frozen.get(key) != live.get(key)
        }
        rows.append(
            {
                "case": name,
                "analytical": "MATCH" if not diffs else "DIFF",
                "analysis_id": data.get("analysis_id"),
                "request_id": envelope.get("request_id"),
                "dung": live.get("overall_dung"),
                "elapsed_s": elapsed,
                "diffs": diffs,
            }
        )
        if diffs:
            mismatches.append({"case": name, "diffs": diffs})
        if name == DUNG:
            dung_payload = data
            dung_spec = spec

    export: dict[str, object] = {}
    if dung_payload and dung_spec:
        report_input = prepare_customer_report_input(
            analysis_id=str(dung_payload.get("analysis_id")),
            source="current",
            data=dung_payload,
            birth_input={**dung_spec, "full_name": DUNG},
        )
        pdf_path, pdf_name, _pdf = export_customer_file(report_input=report_input, fmt="pdf")
        raw = Path(pdf_path).read_bytes()
        export["pdf"] = {
            "name": pdf_name,
            "bytes": len(raw),
            "mime_ok": raw[:5] == b"%PDF-",
        }
        Path(pdf_path).unlink(missing_ok=True)
        docx_path, docx_name, _docx = export_customer_file(report_input=report_input, fmt="docx")
        from docx import Document

        document = Document(str(docx_path))
        blob = "\n".join(p.text for p in document.paragraphs) + "\n".join(
            cell.text for table in document.tables for row in table.rows for cell in row.cells
        )
        export["docx"] = {
            "name": docx_name,
            "bytes": Path(docx_path).stat().st_size,
            "zip_ok": Path(docx_path).read_bytes()[:2] == b"PK",
            "has_dung": "Thủy · Nhâm · Thực Thần" in blob,
            "has_hy": HY in blob,
            "has_dieu_hau": "Điều hậu" in blob or "Hỏa" in blob,
            "has_reason": "TIẾT" in blob or "Tiết" in blob or "tiết" in blob,
        }
        Path(docx_path).unlink(missing_ok=True)

    dung_row = next(r for r in rows if r["case"] == DUNG)
    out = {
        "health": health,
        "api_v1_health": api_health,
        "version": version,
        "mismatch_count": len(mismatches),
        "ten": rows,
        "export": export,
        "identity_ok": bool(dung_row.get("analysis_id")) and dung_row.get("analysis_id") == dung_row.get("request_id"),
        "pass": (
            health.get("status") == "ok"
            and version.get("api_version")
            and not mismatches
            and bool(export.get("pdf", {}).get("mime_ok"))
            and bool(export.get("docx", {}).get("zip_ok"))
            and bool(export.get("docx", {}).get("has_dung"))
        ),
    }
    OUT.write_text(json.dumps(out, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({"pass": out["pass"], "mismatch_count": out["mismatch_count"]}, ensure_ascii=False))
    return 0 if out["pass"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
