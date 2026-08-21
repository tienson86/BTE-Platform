"""G3-02 Linux/runtime smoke: health, version, ten-case truth, PDF/DOCX, locale.

Run from repository root (or /app in the API image):

    python release/gate_03/_g3_02_linux_smoke.py

Does not write engines. Does not change Gate-2 semantics.
Request IDs must be ASCII (httpx header encoding).
"""

from __future__ import annotations

import json
import locale
import os
import platform
import sys
import time
from pathlib import Path
from zoneinfo import ZoneInfo

ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(ROOT))
sys.path.insert(0, str(ROOT / "release" / "gate_02"))

from fastapi.testclient import TestClient  # noqa: E402

from applications.api.app import app  # noqa: E402
from applications.api.services.customer_export import (  # noqa: E402
    export_customer_file,
    prepare_customer_report_input,
)
from _g2_01_binding_probe import CASES, COMPARE_KEYS, FROZEN, fingerprint  # noqa: E402

OUT = Path(__file__).resolve().parent / "G3_02_SMOKE.json"
EXPORT_NAMES = ("Ngô Đắc Dũng", "Vũ Thị Thanh Tuyền")
HY = "Chưa đủ căn cứ xác định Hỷ thần bổ trợ riêng"


def _runtime_facts() -> dict[str, object]:
    tz_ok = True
    tz_error = ""
    try:
        ZoneInfo("Asia/Ho_Chi_Minh")
    except Exception as exc:  # noqa: BLE001 — record, do not hide
        tz_ok = False
        tz_error = str(exc)
    chromium = os.environ.get("PLAYWRIGHT_BROWSERS_PATH", "")
    return {
        "platform": sys.platform,
        "system": platform.system(),
        "python": sys.version.split()[0],
        "executable": sys.executable,
        "locale": locale.getpreferredencoding(False),
        "lang": os.environ.get("LANG", ""),
        "lc_all": os.environ.get("LC_ALL", ""),
        "tz": os.environ.get("TZ", ""),
        "asia_ho_chi_minh": tz_ok,
        "tz_error": tz_error,
        "chromium": chromium,
        "cwd": str(Path.cwd()),
        "repo_root": str(ROOT),
        "database_exists": (ROOT / "database").is_dir(),
        "knowledge_expert_translation": (
            ROOT / "knowledge" / "expert_translation" / "translation_rules.json"
        ).is_file(),
        "luck_package": (
            ROOT / "knowledge" / "packages" / "luck" / "foundation" / "PACKAGE.json"
        ).is_file(),
    }


def _export_one(name: str, payload: dict, spec: dict) -> dict[str, object]:
    from docx import Document

    report_input = prepare_customer_report_input(
        analysis_id=str(payload.get("analysis_id")),
        source="current",
        data=payload,
        birth_input={**spec, "full_name": name},
    )
    pdf_path, pdf_name, _pdf = export_customer_file(report_input=report_input, fmt="pdf")
    pdf_bytes = Path(pdf_path).read_bytes()
    Path(pdf_path).unlink(missing_ok=True)
    docx_path, docx_name, _docx = export_customer_file(report_input=report_input, fmt="docx")
    document = Document(str(docx_path))
    blob = "\n".join(p.text for p in document.paragraphs) + "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    size = Path(docx_path).stat().st_size
    zip_ok = Path(docx_path).read_bytes()[:2] == b"PK"
    Path(docx_path).unlink(missing_ok=True)
    return {
        "case": name,
        "pdf_name": pdf_name,
        "pdf_bytes": len(pdf_bytes),
        "pdf_signature": pdf_bytes[:5] == b"%PDF-",
        "docx_name": docx_name,
        "docx_bytes": size,
        "docx_zip": zip_ok,
        "docx_has_hy": HY in blob,
        "docx_has_vietnamese": all(ch in blob for ch in ("Đ", "ă", "â")),
    }


def main() -> int:
    frozen_by = {row["case"]: row for row in json.loads(FROZEN.read_text(encoding="utf-8"))}
    facts = _runtime_facts()
    client = TestClient(app)
    health = client.get("/health").json()
    version = client.get("/version").json()
    api_health = client.get("/api/v1/health").json()
    portal_ok = False
    try:
        from applications.customer_portal.app import app as portal_app

        portal_ok = TestClient(portal_app).get("/healthz").json().get("status") == "ok"
    except Exception as exc:  # noqa: BLE001
        facts["portal_health_error"] = str(exc)

    mismatches: list[dict[str, object]] = []
    rows: list[dict[str, object]] = []
    payloads: dict[str, tuple[dict, dict]] = {}
    for index, spec in enumerate(CASES):
        name = str(spec["name"])
        body = {k: v for k, v in spec.items() if k != "name"}
        body["full_name"] = name
        request_id = f"g3-02-{index}"
        started = time.perf_counter()
        response = client.post(
            "/api/v1/analyze",
            json=body,
            headers={"X-Request-ID": request_id},
        )
        elapsed = round(time.perf_counter() - started, 3)
        response.raise_for_status()
        envelope = response.json()
        data = envelope["data"]
        live = fingerprint(name, data)
        frozen = frozen_by[name]
        diffs = {
            key: {"frozen": frozen.get(key), "live": live.get(key)}
            for key in COMPARE_KEYS
            if frozen.get(key) != live.get(key)
        }
        rows.append(
            {
                "case": name,
                "analytical": "MATCH" if not diffs else "DIFF",
                "analysis_id": data.get("analysis_id"),
                "request_id": envelope.get("request_id"),
                "elapsed_s": elapsed,
                "diffs": diffs,
            }
        )
        if diffs:
            mismatches.append({"case": name, "diffs": diffs})
        if name in EXPORT_NAMES:
            payloads[name] = (data, spec)

    exports: dict[str, object] = {}
    for name, (payload, spec) in payloads.items():
        exports[name] = _export_one(name, payload, spec)

    result = {
        "runtime": facts,
        "health": health,
        "api_v1_health": api_health,
        "version": version,
        "portal_healthz": portal_ok,
        "mismatch_count": len(mismatches),
        "ten": rows,
        "export": exports,
        "pass": (
            health.get("status") == "ok"
            and version.get("api_version") == "1.0.0"
            and not mismatches
            and facts.get("asia_ho_chi_minh") is True
            and facts.get("database_exists") is True
            and all(item.get("pdf_signature") and item.get("docx_zip") for item in exports.values())
        ),
    }
    OUT.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({"pass": result["pass"], "mismatch_count": result["mismatch_count"], "python": facts["python"], "platform": facts["platform"]}, ensure_ascii=False))
    return 0 if result["pass"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
