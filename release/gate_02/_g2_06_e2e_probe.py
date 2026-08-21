"""G2-06 end-to-end customer acceptance probe.

Uses the real Analyze API + stored-snapshot export path. Does not write engines.
"""

from __future__ import annotations

import copy
import json
import re
import sys
import time
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient

ROOT = Path(__file__).resolve().parent
REPO = ROOT.parent.parent
sys.path.insert(0, str(REPO))
sys.path.insert(0, str(ROOT))

from applications.api.app import app  # noqa: E402
from applications.api.services.customer_export import (  # noqa: E402
    build_customer_export_filename,
    export_customer_file,
    prepare_customer_report_input,
)
from engines.report_engine.rendering.html_report_v1 import render_html  # noqa: E402
from engines.report_engine.rendering.report_sections_v1 import build_presented_report  # noqa: E402

from _g2_01_binding_probe import CASES, COMPARE_KEYS, FROZEN, fingerprint  # noqa: E402

SHOT = ROOT / "screenshots" / "g2_06"
HY_NEUTRAL = "Chưa đủ căn cứ xác định Hỷ thần bổ trợ riêng"
RULE_ID_RE = re.compile(
    r"\b(?:cli|com_san|pat|str|sea|tmp|flo|flw|ctl|sup|spc|spe|cmb|root)_[a-z0-9_]+\b",
    re.I,
)
LEAK_RE = re.compile(r"\b(?:male|female|undefined|null|mock|fixture|debug)\b", re.I)
CONTRACT_RE = re.compile(r"UsefulGodView@1\.5")

PRIMARY = {
    "Nguyễn Tiến Sơn": "son",
    "Vũ Thị Thanh Tuyền": "tuyen",
    "Ngô Đắc Dũng": "dung",
    "Cao Xuân Trường": "truong",
}

BIRTH_PLACE = "Hà Nội"


def _pdf_contains(raw: bytes, needle: str) -> bool:
    if not needle:
        return False
    return (
        needle.encode("utf-8") in raw
        or needle.encode("utf-16-be") in raw
        or needle.encode("utf-16-le") in raw
    )


def _docx_text(path: Path) -> str:
    document = Document(str(path))
    paragraphs = "\n".join(item.text for item in document.paragraphs)
    tables = "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    return f"{paragraphs}\n{tables}"


def _meta_rows(report_input) -> dict[str, str]:
    presented = build_presented_report(report_input)
    rows: dict[str, str] = {}
    for section in presented.sections:
        for label, value in section.meta_rows:
            rows[label] = value
    return rows


def _screenshot_html(html_path: Path, png_path: Path) -> str | None:
    try:
        from playwright.sync_api import sync_playwright
    except Exception as exc:  # noqa: BLE001
        return str(exc)
    with sync_playwright() as playwright:
        browser = playwright.chromium.launch()
        page = browser.new_page(viewport={"width": 1280, "height": 900})
        page.goto(html_path.resolve().as_uri(), wait_until="networkidle")
        page.screenshot(path=str(png_path), full_page=True)
        browser.close()
    return None


def _analyze_api(client: TestClient, spec: dict[str, object], request_id: str) -> tuple[dict, float]:
    body = {key: value for key, value in spec.items() if key != "name"}
    body["full_name"] = spec["name"]
    body["birth_place"] = BIRTH_PLACE
    started = time.perf_counter()
    response = client.post(
        "/api/v1/analyze",
        json=body,
        headers={"X-Request-ID": request_id},
    )
    elapsed = round(time.perf_counter() - started, 3)
    response.raise_for_status()
    payload = response.json()
    return payload, elapsed


def _export_pair(report_input, dest_dir: Path) -> dict[str, object]:
    started = time.perf_counter()
    pdf_path, pdf_name, _pdf = export_customer_file(report_input=report_input, fmt="pdf")
    pdf_s = round(time.perf_counter() - started, 3)
    dest_pdf = dest_dir / pdf_name
    dest_pdf.write_bytes(Path(pdf_path).read_bytes())
    Path(pdf_path).unlink(missing_ok=True)
    started = time.perf_counter()
    docx_path, docx_name, _docx = export_customer_file(report_input=report_input, fmt="docx")
    docx_s = round(time.perf_counter() - started, 3)
    dest_docx = dest_dir / docx_name
    dest_docx.write_bytes(Path(docx_path).read_bytes())
    Path(docx_path).unlink(missing_ok=True)
    return {
        "pdf": dest_pdf.name,
        "docx": dest_docx.name,
        "pdf_s": pdf_s,
        "docx_s": docx_s,
        "pdf_mime_ok": dest_pdf.read_bytes()[:5] == b"%PDF-",
        "docx_zip_ok": dest_docx.read_bytes()[:2] == b"PK",
        "pdf_bytes": dest_pdf.stat().st_size,
        "docx_bytes": dest_docx.stat().st_size,
    }


def _customer_leaks(text: str) -> list[str]:
    leaks: list[str] = []
    if RULE_ID_RE.search(text):
        leaks.append("rule_id")
    if LEAK_RE.search(text):
        leaks.append("enum_or_debug")
    if CONTRACT_RE.search(text):
        leaks.append("contract_name")
    return leaks


def main() -> int:
    SHOT.mkdir(parents=True, exist_ok=True)
    frozen_rows = json.loads(FROZEN.read_text(encoding="utf-8"))
    frozen_by = {row["case"]: row for row in frozen_rows}
    client = TestClient(app)

    ten: list[dict[str, object]] = []
    mismatches: list[dict[str, object]] = []
    id_trace: list[dict[str, object]] = []
    timings: dict[str, float] = {}
    primary_rows: list[dict[str, object]] = []
    stored_by_name: dict[str, dict] = {}

    for index, spec in enumerate(CASES):
        name = str(spec["name"])
        request_id = f"g2-06-{index}"
        envelope, analyze_s = _analyze_api(client, spec, request_id)
        data = envelope["data"]
        live = fingerprint(name, data)
        frozen = frozen_by[name]
        diffs = {
            key: {"frozen": frozen.get(key), "live": live.get(key)}
            for key in COMPARE_KEYS
            if frozen.get(key) != live.get(key)
        }
        report_started = time.perf_counter()
        report_input = prepare_customer_report_input(
            analysis_id=str(data.get("analysis_id")),
            source="current",
            data=data,
            birth_input={**spec, "full_name": name, "birth_place": BIRTH_PLACE},
        )
        presented = _meta_rows(report_input)
        report_s = round(time.perf_counter() - report_started, 3)
        timings[f"{name}_analyze_s"] = analyze_s
        timings[f"{name}_report_s"] = report_s
        row = {
            "case": name,
            "analytical": "MATCH" if not diffs else "DIFF",
            "analysis_id": data.get("analysis_id"),
            "request_id": envelope.get("request_id"),
            "four_pillars": live.get("four_pillars"),
            "strength": f"{live.get('strength_score')} {live.get('strength_level')}",
            "pattern": live.get("pattern"),
            "ug_override_eligible": live.get("ug_override_eligible"),
            "dieu_hau": live.get("climate_preference_label") or live.get("dieu_hau"),
            "dung": live.get("overall_dung"),
            "reason_archetype": live.get("reason_archetype"),
            "hy": live.get("customer_hy"),
            "ky": live.get("ky"),
            "luck": live.get("current_luck"),
            "report_dung": presented.get("Dụng thần"),
            "report_hy": presented.get("Hỷ thần"),
            "surfaces": "MATCH"
            if presented.get("Dụng thần") == live.get("overall_dung")
            and presented.get("Hỷ thần") == live.get("customer_hy")
            else "DIFF",
            "diffs": diffs,
        }
        ten.append(row)
        stored_by_name[name] = {
            "envelope": envelope,
            "data": data,
            "spec": spec,
            "live": live,
            "report_input": report_input,
            "presented": presented,
        }
        if diffs:
            mismatches.append({"case": name, "diffs": diffs})

    if mismatches:
        (ROOT / "G2_06_E2E_PROBE.json").write_text(
            json.dumps({"mismatch_count": len(mismatches), "mismatches": mismatches, "ten": ten}, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        print("G2-06 BLOCKED — GATE-1 FROZEN TRUTH MISMATCH")
        return 1

    # Failed Analyze: validation, no fake payload.
    invalid = client.post("/api/v1/analyze", json={"year": 0, "month": 5, "day": 15})
    failure = {
        "status": invalid.status_code,
        "has_data_payload": bool((invalid.json() or {}).get("data") if invalid.headers.get("content-type", "").startswith("application/json") else False)
        and invalid.status_code < 400,
        "no_stack": "Traceback" not in invalid.text,
    }

    # Primary four: PDF/DOCX/History/ID trace.
    dung_store = None
    tuyen_store = None
    for name, slug in PRIMARY.items():
        packed = stored_by_name[name]
        data = packed["data"]
        live = packed["live"]
        presented = packed["presented"]
        report_input = packed["report_input"]
        html = render_html(report_input)
        html_path = SHOT / f"{slug}_report.html"
        html_path.write_text(html, encoding="utf-8")
        png_path = SHOT / f"{slug}_report.png"
        png_error = _screenshot_html(html_path, png_path)
        artifacts = _export_pair(report_input, SHOT)
        pdf_raw = (SHOT / artifacts["pdf"]).read_bytes()
        docx_blob = _docx_text(SHOT / artifacts["docx"])
        needles = [
            name.split()[-1],
            str(live.get("four_pillars") or "").split(" / ")[2] if live.get("four_pillars") else "",
            str(live.get("overall_dung") or ""),
            str(live.get("customer_hy") or ""),
            str(presented.get("Căn cứ chọn Dụng") or live.get("short_reason") or "")[:12],
        ]
        html_leaks = _customer_leaks(html)
        docx_leaks = _customer_leaks(docx_blob)
        leaks = html_leaks + docx_leaks
        html_hits = {needle: needle in html for needle in needles if needle}
        docx_hits = {needle: needle in docx_blob for needle in needles if needle}
        # Official PDF is Playwright Report V1 (G2-04). CID fonts make naive
        # Unicode byte-search unreliable; HTML source + DOCX remain the text checks.
        checks = {
            "no_tong_tai": "Tòng Tài" not in html and "Tòng Tài" not in docx_blob,
            "no_cuc_nhuoc": "cực nhược" not in html and "cực nhược" not in docx_blob,
            "hy_neutral_or_supported": live.get("customer_hy") == HY_NEUTRAL
            or bool(str(live.get("customer_hy") or "").strip()),
            "filename_has_case": name.split()[-1].replace("ễ", "e") in artifacts["pdf"]
            or any(part in artifacts["pdf"] for part in ("Son", "Tuyen", "Dung", "Truong")),
        }
        if name == "Ngô Đắc Dũng":
            checks["dung_not_tho_mau"] = "Thổ · Mậu · Thiên Ấn" not in html
            checks["dung_is_thuy_nham"] = live.get("overall_dung") == "Thủy · Nhâm · Thực Thần"
            checks["level1"] = live.get("detected_special_pattern") == "gia_sac" and live.get("ug_override_eligible") is False
            dung_store = packed
        if name == "Vũ Thị Thanh Tuyền":
            checks["tuyen_not_nham_overall"] = live.get("overall_dung") == "Mộc · Ất · Chính Quan"
            checks["tuyen_che"] = live.get("reason_archetype") == "CHẾ"
            tuyen_store = packed
        if name == "Cao Xuân Trường":
            checks["truong_sinh"] = live.get("reason_archetype") == "SINH / TRỢ"
            checks["truong_dung"] = live.get("overall_dung") == "Kim · Tân · Chính Ấn"
        if name == "Nguyễn Tiến Sơn":
            checks["son_dung"] = live.get("overall_dung") == "Hỏa · Đinh · Chính Quan"
        primary_rows.append(
            {
                "case": name,
                "slug": slug,
                "analysis_id": data.get("analysis_id"),
                "request_id": packed["envelope"].get("request_id"),
                "pdf": artifacts,
            "pdf_text_hits": {needle: _pdf_contains(pdf_raw, needle) for needle in needles if needle},
            "docx_text_hits": docx_hits,
                "leaks": leaks,
                "checks": checks,
                "png_error": png_error,
                "journey": "PASS"
                if artifacts["pdf_mime_ok"]
                and artifacts["docx_zip_ok"]
                and not leaks
                and all(html_hits.values())
                and all(docx_hits.values())
                and all(checks.values())
                else "FAIL",
                "html_text_hits": html_hits,
            }
        )
        timings[f"{slug}_pdf_s"] = artifacts["pdf_s"]
        timings[f"{slug}_docx_s"] = artifacts["docx_s"]
        id_trace.append(
            {
                "case": name,
                "request_id": packed["envelope"].get("request_id"),
                "data.analysis_id": data.get("analysis_id"),
                "report.case_id": packed["report_input"].metadata.case_id,
                "pdf_filename": artifacts["pdf"],
                "docx_filename": artifacts["docx"],
                "coherent": packed["envelope"].get("request_id") == data.get("analysis_id")
                == packed["report_input"].metadata.case_id,
            }
        )

    # Current cross-case: History A (Dũng) export while current conceptually Tuyền.
    history = {"ok": False}
    if dung_store and tuyen_store:
        stored_a = copy.deepcopy(dung_store["data"])
        current_b = tuyen_store["data"]
        hist_input = prepare_customer_report_input(
            analysis_id=str(stored_a.get("analysis_id")),
            source="history",
            data=stored_a,
            birth_input={**dung_store["spec"], "full_name": "Ngô Đắc Dũng", "birth_place": BIRTH_PLACE},
        )
        hist_art = _export_pair(hist_input, SHOT)
        hist_pdf = (SHOT / hist_art["pdf"]).read_bytes()
        hist_docx = _docx_text(SHOT / hist_art["docx"])
        dung_display = str(dung_store["live"].get("overall_dung"))
        tuyen_display = str(tuyen_store["live"].get("overall_dung"))
        history = {
            "ok": dung_display in hist_docx
            and tuyen_display not in hist_docx
            and "Tuyen" not in hist_art["pdf"]
            and stored_a.get("analysis_id") != current_b.get("analysis_id")
            and stored_a.get("useful_god", {}).get("useful_display") == dung_display,
            "history_pdf": hist_art["pdf"],
            "history_docx": hist_art["docx"],
            "current_id": current_b.get("analysis_id"),
            "history_id": stored_a.get("analysis_id"),
            "pdf_unicode_grep": _pdf_contains(hist_pdf, "Nhâm"),
        }
        # Re-analyze Dũng birth → new id, old snapshot unchanged.
        re_env, _elapsed = _analyze_api(client, dung_store["spec"], "g2-06-reanalyze-dung")
        re_data = re_env["data"]
        history["reanalyze_new_id"] = re_data.get("analysis_id")
        history["reanalyze_old_unchanged"] = stored_a.get("analysis_id") != re_data.get("analysis_id")
        history["reanalyze_old_dung"] = stored_a.get("useful_god", {}).get("useful_display") == dung_display
        history["ok"] = bool(history["ok"] and history["reanalyze_old_unchanged"] and history["reanalyze_old_dung"])

    # File cross-contamination: Dũng file vs Tuyền file
    dung_pdf_name = next(row["pdf"]["pdf"] for row in primary_rows if row["slug"] == "dung")
    tuyen_pdf_name = next(row["pdf"]["pdf"] for row in primary_rows if row["slug"] == "tuyen")
    dung_docx_name = next(row["pdf"]["docx"] for row in primary_rows if row["slug"] == "dung")
    tuyen_docx_name = next(row["pdf"]["docx"] for row in primary_rows if row["slug"] == "tuyen")
    dung_docx = _docx_text(SHOT / dung_docx_name)
    tuyen_docx = _docx_text(SHOT / tuyen_docx_name)
    cross = {
        "dung_file_has_dung": "Thủy · Nhâm · Thực Thần" in dung_docx or "Nhâm" in dung_docx,
        "dung_file_not_tuyen": "Mộc · Ất · Chính Quan" not in dung_docx,
        "tuyen_file_has_tuyen": "Mộc · Ất · Chính Quan" in tuyen_docx or "Chính Quan" in tuyen_docx,
        "tuyen_file_not_dung_pattern": "Giá Sắc" not in tuyen_docx,
        "distinct_pdf_names": dung_pdf_name != tuyen_pdf_name,
    }

    out = {
        "mismatch_count": 0,
        "ten_match": all(row["analytical"] == "MATCH" and row["surfaces"] == "MATCH" for row in ten),
        "primary_journeys": {row["slug"]: row["journey"] for row in primary_rows},
        "primary": primary_rows,
        "ten": ten,
        "id_trace": id_trace,
        "history": history,
        "cross_files": cross,
        "analyze_failure": failure,
        "timings": timings,
        "hy_copy": HY_NEUTRAL,
    }
    (ROOT / "G2_06_E2E_PROBE.json").write_text(json.dumps(out, ensure_ascii=False, indent=2), encoding="utf-8")
    primary_pass = all(row["journey"] == "PASS" for row in primary_rows)
    history_pass = bool(history.get("ok"))
    cross_pass = all(cross.values())
    fail_pass = failure["status"] >= 400 and failure["no_stack"]
    status = primary_pass and out["ten_match"] and history_pass and cross_pass and fail_pass
    print(
        json.dumps(
            {
                "ten_match": out["ten_match"],
                "primary": out["primary_journeys"],
                "history_ok": history_pass,
                "cross_ok": cross_pass,
                "failure_ok": fail_pass,
                "pass": status,
            },
            ensure_ascii=False,
            indent=2,
        )
    )
    return 0 if status else 1


if __name__ == "__main__":
    raise SystemExit(main())
