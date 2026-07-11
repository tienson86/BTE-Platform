"""
report_engine/word_renderer.py

Word Renderer.

Nhiệm vụ:
- Xuất báo cáo Bát Tự dạng DOCX.
- Tạo tài liệu có thể chỉnh sửa.

Sử dụng:
python-docx
"""


from __future__ import annotations


from docx import Document

from docx.shared import Pt




class WordRenderer:
    """
    Bộ xuất file Word.
    """

    name = "WordRenderer"



    def __init__(self):

        self.document = None



    def render(
        self,
        report: dict,
        output_file: str
    ):
        """
        Tạo file Word.

        Parameters
        ----------
        report:
            Report Object

        output_file:
            đường dẫn docx
        """


        self.document = Document()



        self.setup_style()



        self.add_title()



        self.add_person_info(
            report
        )



        self.add_score(
            report
        )



        self.add_summary(
            report
        )



        self.add_sections(
            report
        )



        self.document.save(
            output_file
        )




    def setup_style(self):
        """
        Cấu hình font mặc định.
        """

        style = (
            self.document.styles["Normal"]
        )


        style.font.name = "Arial"

        style.font.size = Pt(11)




    def add_title(self):


        title = (
            self.document
            .add_heading(
                "BÁO CÁO PHÂN TÍCH BÁT TỰ",
                level=1
            )
        )



    def add_person_info(
        self,
        report
    ):


        self.document.add_heading(

            "Thông Tin Khách Hàng",

            level=2

        )


        person = report.get(
            "person",
            {}
        )


        self.document.add_paragraph(

            f"""
Họ tên:
{person.get('name','')}

Ngày sinh:
{person.get('birth_date','')}

Giờ sinh:
{person.get('birth_time','')}

Giới tính:
{person.get('gender','')}
"""

        )




    def add_score(
        self,
        report
    ):


        self.document.add_heading(

            "Đánh Giá Tổng Quan",

            level=2

        )


        score = report.get(
            "score",
            {}
        )


        self.document.add_paragraph(

            f"""
Tổng điểm:
{score.get('overall',0)}

Xếp loại:
{score.get('rating','')}
"""

        )




    def add_summary(
        self,
        report
    ):


        summary = report.get(
            "summary",
            {}
        )


        self.document.add_heading(

            "Tổng Kết Mệnh Cục",

            level=2

        )


        self.document.add_heading(

            "Điểm mạnh",

            level=3

        )


        for item in summary.get(
            "strengths",
            []
        ):


            self.document.add_paragraph(

                item.get(
                    "title",
                    ""
                ),

                style="List Bullet"

            )



        self.document.add_heading(

            "Điểm cần lưu ý",

            level=3

        )


        for item in summary.get(
            "weaknesses",
            []
        ):


            self.document.add_paragraph(

                item.get(
                    "title",
                    ""
                ),

                style="List Bullet"

            )




    def add_sections(
        self,
        report
    ):


        self.document.add_heading(

            "Phân Tích Chi Tiết",

            level=2

        )



        for section in report.get(
            "sections",
            []
        ):


            self.document.add_heading(

                section.get(
                    "title",
                    ""
                ),

                level=3

            )



            for item in section.get(
                "items",
                []
            ):


                self.document.add_heading(

                    item.get(
                        "title",
                        ""
                    ),

                    level=4

                )


                self.document.add_paragraph(

                    f"""
Điểm đánh giá:
{item.get('score',0)}

{item.get('content','')}
"""

                )
