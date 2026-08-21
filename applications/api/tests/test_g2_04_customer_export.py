"""G2-04 customer export: identity, contract, presentation, DOCX tables, no legacy fallback."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient

from applications.api.app import app
from applications.api.exceptions import CustomerExportError
from applications.api.services.customer_contract import (
    CONTRACT_INCOMPLETE_MESSAGE,
    CONTRACT_MISMATCH_MESSAGE,
    EMPTY_RESULT_MESSAGE,
    HISTORY_MISMATCH_MESSAGE,
)
from applications.api.services.customer_export import (
    build_customer_export_filename,
    export_customer_file,
    prepare_customer_report_input,
)
from applications.api.services.orchestrator import OrchestratorService
from applications.api.services.result_identity import (
    CUSTOMER_USEFUL_GOD_CONTRACT,
    stamp_customer_result_identity,
)
from engines.report_engine.rendering.html_report_v1 import render_html
from engines.report_engine.rendering.report_sections_v1 import build_presented_report
from engines.useful_god_engine.presentation import INSUFFICIENT_CUSTOMER_FAVORABLE_DISPLAY

DUNG = {
    "name": "Ngô Đắc Dũng",
    "year": 1985,
    "month": 9,
    "day": 18,
    "hour": 8,
    "minute": 0,
    "gender": "male",
    "timezone": "Asia/Bangkok",
}
TUYEN = {
    "name": "Vũ Thị Thanh Tuyền",
    "year": 1984,
    "month": 7,
    "day": 13,
    "hour": 21,
    "minute": 1,
    "gender": "female",
    "timezone": "Asia/Bangkok",
}

HY_NEUTRAL = INSUFFICIENT_CUSTOMER_FAVORABLE_DISPLAY
DUNG_DISPLAY = "Thủy · Nhâm · Thực Thần"
TUYEN_DISPLAY = "Mộc · Ất · Chính Quan"


def _analyze(spec: dict[str, object], analysis_id: str) -> dict:
    kwargs = {key: value for key, value in spec.items() if key != "name"}
    payload = OrchestratorService().analyze(**kwargs)
    stamped = stamp_customer_result_identity(payload, analysis_id)
    return stamped


def _docx_blob(path: Path) -> str:
    document = Document(str(path))
    paragraphs = "\n".join(item.text for item in document.paragraphs)
    tables = "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    return f"{paragraphs}\n{tables}"


def test_missing_result_is_blocked() -> None:
    with pytest.raises(CustomerExportError) as caught:
        prepare_customer_report_input(analysis_id="x", source="current", data={})
    assert caught.value.code == "export_missing_result"
    assert EMPTY_RESULT_MESSAGE in caught.value.message


def test_contract_mismatch_blocks_official_export() -> None:
    with pytest.raises(CustomerExportError) as caught:
        prepare_customer_report_input(
            analysis_id="old-1",
            source="history",
            data={"analysis_id": "old-1", "pattern": {"cach_cuc": "Chính Ấn"}},
        )
    assert caught.value.status_code == 409
    assert caught.value.code == "export_contract_mismatch"
    assert CONTRACT_MISMATCH_MESSAGE in caught.value.message


def test_history_mismatch_does_not_export_current() -> None:
    with pytest.raises(CustomerExportError) as caught:
        prepare_customer_report_input(
            analysis_id="hist-a",
            source="history",
            data={
                "analysis_id": "current-b",
                "useful_god_source": {"contract": CUSTOMER_USEFUL_GOD_CONTRACT},
                "useful_god": {"useful_display": DUNG_DISPLAY, "favorable_display": HY_NEUTRAL},
            },
        )
    assert caught.value.code == "export_history_mismatch"
    assert HISTORY_MISMATCH_MESSAGE in caught.value.message


def test_dung_presentation_retains_analysis_id_and_hk_r1h() -> None:
    payload = _analyze(DUNG, "g2-04-dung")
    report_input = prepare_customer_report_input(
        analysis_id="g2-04-dung",
        source="current",
        data=payload,
        birth_input=DUNG,
    )
    assert report_input.metadata.case_id == "g2-04-dung"
    assert report_input.useful_god.useful_display == DUNG_DISPLAY
    assert report_input.useful_god.favorable_display == HY_NEUTRAL
    assert report_input.useful_god.short_reason
    assert "TIẾT" in (report_input.useful_god.reason_archetype or report_input.useful_god.short_reason)
    assert report_input.useful_god.climate_preference_label
    assert "Hỏa" in report_input.useful_god.climate_preference_label
    assert report_input.useful_god.useful_display != report_input.useful_god.climate_display
    assert report_input.pattern.primary_pattern
    assert "Chuyên cách ưu tiên Ấn" not in (report_input.pattern.explanation or "")
    presented = build_presented_report(report_input)
    blob = "\n".join(
        f"{label}: {value}"
        for section in presented.sections
        for label, value in section.meta_rows
    )
    assert DUNG_DISPLAY in blob
    assert HY_NEUTRAL in blob
    assert "Căn cứ chọn Dụng" in blob
    assert "g2-04-dung" in presented.footer
    html = render_html(report_input)
    assert DUNG_DISPLAY in html
    assert HY_NEUTRAL in html
    assert "Phân bố Ngũ hành" in html
    assert "không phải mức vượng suy" in html
    assert "1.00" in html
    filename = build_customer_export_filename(report_input, "pdf")
    assert filename.startswith("BTE_BaoCao_")
    assert "NgoDacDung" in filename or "Dung" in filename
    assert filename.endswith(".pdf")
    assert "g2-04-dung.pdf" != filename


def test_tuyen_docx_tables_and_paragraphs(tmp_path: Path) -> None:
    payload = _analyze(TUYEN, "g2-04-tuyen")
    report_input = prepare_customer_report_input(
        analysis_id="g2-04-tuyen",
        source="history",
        data=payload,
        birth_input=TUYEN,
    )
    path, name, result = export_customer_file(
        report_input=report_input,
        fmt="docx",
        service=None,
    )
    text = _docx_blob(Path(path))
    assert TUYEN_DISPLAY in text
    assert HY_NEUTRAL in text
    assert "Tòng Tài" not in text
    assert "cực nhược" not in text
    assert "0.66" in text
    assert "Kiếp Tài" in text
    assert "g2-04-tuyen" in text
    assert name.endswith(".docx")
    assert result.case_id == "g2-04-tuyen"
    Path(path).unlink(missing_ok=True)


def test_api_docx_route_sets_mime_and_analysis_header() -> None:
    payload = _analyze(DUNG, "g2-04-api-dung")
    client = TestClient(app)
    response = client.post(
        "/api/v1/export/docx",
        json={
            "analysis_id": "g2-04-api-dung",
            "source": "current",
            "data": payload,
            "input": DUNG,
        },
    )
    assert response.status_code == 200
    assert "officedocument.wordprocessingml.document" in response.headers["content-type"]
    assert response.headers["X-BTE-Analysis-Id"] == "g2-04-api-dung"
    assert "attachment" in response.headers["content-disposition"]
    assert response.content[:2] == b"PK"
    assert len(response.content) > 2048


def test_api_mismatch_returns_customer_notice_not_file() -> None:
    client = TestClient(app)
    response = client.post(
        "/api/v1/export/pdf",
        json={
            "analysis_id": "stale",
            "source": "history",
            "data": {"analysis_id": "stale", "bazi": {"day_master": "Canh"}},
        },
    )
    assert response.status_code == 409
    body = response.json()
    assert body["success"] is False
    assert CONTRACT_MISMATCH_MESSAGE in body["message"]
    assert "stack" not in str(body).lower()


def test_official_pdf_is_searchable_vietnamese(tmp_path: Path) -> None:
    payload = _analyze(DUNG, "g2-04-pdf-dung")
    report_input = prepare_customer_report_input(
        analysis_id="g2-04-pdf-dung",
        source="current",
        data=payload,
        birth_input=DUNG,
    )
    path, _name, _result = export_customer_file(report_input=report_input, fmt="pdf")
    raw = Path(path).read_bytes()
    extracted = _extract_pdf_text(raw)
    html = render_html(report_input)
    assert raw.startswith(b"%PDF")
    assert b"/Font" in raw
    assert b"/ToUnicode" in raw or b"/ActualText" in raw or b"/MarkInfo" in raw
    assert "Ngô Đắc Dũng" in extracted
    assert DUNG_DISPLAY in html
    assert HY_NEUTRAL in html
    Path(path).unlink(missing_ok=True)


def _extract_pdf_text(raw: bytes) -> str:
    import re
    import zlib

    chunks: list[str] = []
    for match in re.finditer(rb"<FEFF([0-9A-Fa-f]+)>", raw):
        try:
            chunks.append(bytes.fromhex(match.group(1).decode("ascii")).decode("utf-16-be"))
        except (ValueError, UnicodeDecodeError):
            continue
    for match in re.finditer(rb"stream\r?\n(.+?)\r?\nendstream", raw, re.DOTALL):
        payload = match.group(1)
        for candidate in (payload, _inflate(payload)):
            if not candidate:
                continue
            chunks.append(candidate.decode("utf-8", "ignore"))
            chunks.append(candidate.decode("utf-16-be", "ignore"))
            chunks.append(candidate.decode("utf-16-le", "ignore"))
    return "\n".join(chunks)


def _inflate(payload: bytes) -> bytes:
    import zlib

    for wbits in (zlib.MAX_WBITS, -zlib.MAX_WBITS):
        try:
            return zlib.decompress(payload, wbits)
        except zlib.error:
            continue
    return b""


def test_incomplete_payload_message() -> None:
    with pytest.raises(CustomerExportError) as caught:
        prepare_customer_report_input(
            analysis_id="incomplete",
            source="current",
            data={
                "analysis_id": "incomplete",
                "useful_god_source": {"contract": CUSTOMER_USEFUL_GOD_CONTRACT},
                "useful_god": {"useful_display": "", "favorable_display": ""},
            },
        )
    assert CONTRACT_INCOMPLETE_MESSAGE in caught.value.message
