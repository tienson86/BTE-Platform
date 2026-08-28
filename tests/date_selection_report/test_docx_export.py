"""PACK 06 P6-04 Date Selection DOCX export tests."""

from __future__ import annotations

import zipfile
from pathlib import Path

from docx import Document

from engines.date_selection_report.exporting.docx_exporter import (
    DateSelectionDocxExporter,
    extract_docx_text,
)
from engines.date_selection_report.exporting.filename import build_docx_filename
from engines.date_selection_report.exporting.html_projection import project_render_tree_to_html
from engines.date_selection_report.models import DateSelectionReportModel
from engines.date_selection_report.rendering.tree import build_render_tree, create_render_context
from engines.report_engine.contracts.report_export_result_v1 import MEDIA_TYPE_DOCX
from engines.report_engine.exporting.docx_exporter_v1 import validate_docx_file


def _export(presentation_model: DateSelectionReportModel, tmp_path: Path):
    tree = build_render_tree(create_render_context(presentation_model))
    before = presentation_model.to_dict()
    result = DateSelectionDocxExporter().export(tree, tmp_path)
    return tree, result, before, presentation_model.to_dict()


def test_docx_filename_pattern(presentation_model: DateSelectionReportModel) -> None:
    tree = build_render_tree(create_render_context(presentation_model))
    assert build_docx_filename(tree) == "bao-cao-chon-ngay-tot_nguyen-tien-son_09-2026.docx"


def test_docx_generated_and_editable(
    presentation_model: DateSelectionReportModel,
    tmp_path: Path,
) -> None:
    tree, result, before, after = _export(presentation_model, tmp_path)
    output = Path(result.file_path)
    validate_docx_file(output)
    assert result.media_type == MEDIA_TYPE_DOCX
    assert result.format == "docx"
    assert result.file_name.endswith(".docx")
    assert before == after
    document = Document(str(output))
    assert document.paragraphs
    assert len(document.inline_shapes) == 0
    document.paragraphs[0].text = "edited"
    round_trip = tmp_path / "edited.docx"
    document.save(str(round_trip))
    assert "edited" in extract_docx_text(round_trip)


def test_docx_unicode_and_sections(
    presentation_model: DateSelectionReportModel,
    tmp_path: Path,
) -> None:
    tree, result, _, _ = _export(presentation_model, tmp_path)
    text = extract_docx_text(Path(result.file_path))
    assert "Nguyễn Tiến Sơn" in text
    assert "BÁO CÁO CHỌN NGÀY TỐT" in text
    assert "BTE Platform" in text
    assert "Khách hàng" in text
    assert "THÔNG TIN NGƯỜI XEM" in text
    assert "Họ và tên" in text
    assert "Giới tính" in text
    assert "Ngày sinh dương" in text
    assert "Ngày sinh âm" in text
    assert "Can Chi năm" in text
    assert "Nạp âm" in text
    assert "Cung Phi" in text
    assert "Nhóm Trạch" in text
    assert "THÔNG TIN TÌM NGÀY TỐT" in text
    assert "09/2026" in text
    assert "CÁC NGÀY ĐỀ XUẤT" in text
    assert "04/09/2026" in text
    assert "Đại An" in text
    assert "Cấn (Thổ)" in text
    assert "Giờ phù hợp Nhóm Trạch của bạn" in text
    assert "Giờ Thìn (07:01–09:00) · Càn (Kim)" in text
    assert "Tốc Hỷ" in text
    assert "Tiểu Cát" in text
    assert "HƯỚNG DẪN THAM KHẢO" in text
    assert tree.footer.generator in text
    assert "{{" not in text
    assert "Kết quả giờ" not in text
    assert "Lưu Liên" not in text


def test_docx_recommendation_order_and_pdf_truth(
    presentation_model: DateSelectionReportModel,
    tmp_path: Path,
) -> None:
    tree, result, _, _ = _export(presentation_model, tmp_path)
    text = extract_docx_text(Path(result.file_path))
    html = project_render_tree_to_html(tree)
    rec = tree.recommendations[0]
    assert text.index(rec.date_header.solar_date) < text.index(rec.compatible_hours.rows[0].display)
    for hour in rec.compatible_hours.rows:
        assert hour.display in text
        assert hour.display in html
    for group in rec.positive_times.groups:
        assert group.label in text
        assert group.label in html
    named = {style.name for style in Document(result.file_path).styles}
    for style_name in (
        "ReportTitle",
        "SectionTitle",
        "RecommendationTitle",
        "Result",
        "Label",
        "Value",
        "Footer",
    ):
        assert style_name in named


def test_docx_is_openxml_zip(
    presentation_model: DateSelectionReportModel,
    tmp_path: Path,
) -> None:
    _, result, _, _ = _export(presentation_model, tmp_path)
    with zipfile.ZipFile(result.file_path) as archive:
        names = archive.namelist()
    assert "word/document.xml" in names
    assert "[Content_Types].xml" in names
