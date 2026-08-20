"""G1-08 Report/PDF/DOCX bind canonical luck cycles and drop false missing note."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document

from applications.api.services.orchestrator import OrchestratorService
from engines.report_engine.adapters.report_input_v1_adapter import ReportInputV1Adapter
from engines.report_engine.exporting.docx_exporter_v1 import validate_docx_file
from engines.report_engine.localization.labels_vi import FULL_LUCK_CYCLES_GAP_NOTE
from engines.report_engine.rendering.html_report_v1 import render_html
from engines.report_engine.rendering.report_sections_v1 import build_presented_report
from engines.report_engine.services.report_export_service_v1 import ReportExportServiceV1
from tests.report_engine.case_0001_runtime import build_case_0001_source


def test_case_0001_report_luck_has_cycles_without_false_gap_note() -> None:
    report_input = ReportInputV1Adapter().build(build_case_0001_source())
    luck = report_input.luck_cycles
    assert luck.start_age == 5
    assert luck.direction == "forward"
    assert luck.evidence == "Nam · Niên can Bính Dương · Thuận"
    assert luck.method_note == "Khởi vận theo ngày lịch và Tiết khí, độ chính xác theo năm"
    assert luck.precision == "year_level"
    assert luck.cycles[0].summary == "Nhâm Dần"
    assert luck.cycles[3].summary == "Ất Tỵ"
    assert luck.current_gan_zhi == "Ất Tỵ"
    assert luck.current_year_start == 2022
    presented = build_presented_report(report_input)
    section = next(item for item in presented.sections if item.id == "luck-cycles")
    assert FULL_LUCK_CYCLES_GAP_NOTE not in section.notes
    meta = dict(section.meta_rows)
    assert meta["Chiều vận"] == "Thuận"
    assert meta["Tuổi khởi vận"] == "5"
    assert meta["Căn cứ"] == "Nam · Niên can Bính Dương · Thuận"
    assert "độ chính xác theo năm" in meta["Phương pháp V1.0"]
    assert "Ất Tỵ" in meta["Đại vận hiện tại"]
    assert "2022–2031" in meta["Đại vận hiện tại"]
    at_ty = [row for row in section.table.rows if "Ất" in row[1] and "Tỵ" in row[1]][0]
    assert "Ất · Mộc" in at_ty[-1]
    assert "Tỵ · Hỏa" in at_ty[-1]


def test_missing_cycles_still_show_gap_note() -> None:
    from engines.report_engine.contracts.report_input_v1 import (
        ReportInputV1,
        ReportLuckCyclesV1,
        ReportMetadataV1,
        ReportProfileV1,
    )

    presented = build_presented_report(
        ReportInputV1(
            metadata=ReportMetadataV1(case_id="empty-luck"),
            profile=ReportProfileV1(full_name="Test"),
            luck_cycles=ReportLuckCyclesV1(),
        )
    )
    section = next(item for item in presented.sections if item.id == "luck-cycles")
    assert FULL_LUCK_CYCLES_GAP_NOTE in section.notes


def test_api_and_report_same_sequence() -> None:
    payload = OrchestratorService().analyze(
        year=1987, month=1, day=21, hour=4, minute=30, gender="male"
    )
    report = ReportInputV1Adapter().build(build_case_0001_source()).luck_cycles
    luck = payload["luck"]
    assert luck["start_age"] == report.start_age == 5
    assert luck["direction"] == report.direction == "forward"
    assert [item["gan_zhi"] for item in luck["cycles"]] == [
        item.summary for item in report.cycles
    ]
    assert luck["current_cycle"]["gan_zhi"] == report.current_gan_zhi == "Ất Tỵ"


def test_html_and_docx_use_canonical_sequence(tmp_path: Path) -> None:
    report_input = ReportInputV1Adapter().build(build_case_0001_source())
    html = render_html(report_input)
    assert FULL_LUCK_CYCLES_GAP_NOTE not in html
    assert "Nhâm Dần" in html
    assert "Ất Tỵ" in html
    assert "Ất · Mộc" in html
    assert "Tỵ · Hỏa" in html
    assert "Thuận" in html
    assert "Nam · Niên can Bính Dương · Thuận" in html
    result = ReportExportServiceV1(export_root=tmp_path).export_docx(report_input)
    validate_docx_file(Path(result.file_path))
    document = Document(result.file_path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    tables = "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    blob = text + "\n" + tables
    assert "Nhâm Dần" in blob
    assert "Ất Tỵ" in blob
    assert FULL_LUCK_CYCLES_GAP_NOTE not in blob


_ENGLISH_GENDER = re.compile(r"(?i)(?<![a-z])(male|female)(?![a-z])")


def _assert_no_english_gender(blob: str) -> None:
    match = _ENGLISH_GENDER.search(blob)
    assert match is None, f"leaked English gender label: {match.group(0)!r}"


def test_case_0001_report_displays_gioi_tinh_nam(tmp_path: Path) -> None:
    report_input = ReportInputV1Adapter().build(build_case_0001_source())
    assert report_input.profile.gender == "male"
    presented = build_presented_report(report_input)
    chart = next(item for item in presented.sections if item.id == "chart-info")
    assert dict(chart.meta_rows)["Giới tính"] == "Nam"
    html = render_html(report_input)
    assert "Giới tính" in html
    assert "Nam" in html
    _assert_no_english_gender(html)
    result = ReportExportServiceV1(export_root=tmp_path).export_docx(report_input)
    document = Document(result.file_path)
    blob = "\n".join(paragraph.text for paragraph in document.paragraphs)
    tables = "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    combined = blob + "\n" + tables
    assert "Giới tính" in combined
    assert "Nam" in combined
    _assert_no_english_gender(combined)


def test_report_female_displays_nu() -> None:
    from engines.report_engine.contracts.report_input_v1 import (
        ReportInputV1,
        ReportMetadataV1,
        ReportProfileV1,
    )

    presented = build_presented_report(
        ReportInputV1(
            metadata=ReportMetadataV1(case_id="gender-female"),
            profile=ReportProfileV1(full_name="Test", gender="female"),
        )
    )
    chart = next(item for item in presented.sections if item.id == "chart-info")
    assert dict(chart.meta_rows)["Giới tính"] == "Nữ"
    html = render_html(
        ReportInputV1(
            metadata=ReportMetadataV1(case_id="gender-female"),
            profile=ReportProfileV1(full_name="Test", gender="female"),
        )
    )
    assert "Nữ" in html
    _assert_no_english_gender(html)
