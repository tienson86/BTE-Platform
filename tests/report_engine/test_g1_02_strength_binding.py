"""G1-02 Report/PDF/DOCX bind canonical Strength, not Score Engine."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from engines.report_engine.adapters.report_input_v1_adapter import ReportInputV1Adapter
from engines.report_engine.exporting.docx_exporter_v1 import validate_docx_file
from engines.report_engine.rendering.html_report_v1 import render_html
from engines.report_engine.services.report_export_service_v1 import ReportExportServiceV1
from tests.report_engine.case_0001_runtime import build_case_0001_source


def test_case_0001_report_input_uses_canonical_strength() -> None:
    report_input = ReportInputV1Adapter().build(build_case_0001_source())
    assert abs(float(report_input.strength.score or 0.0) - 0.87) < 0.001
    assert report_input.strength.level == "strong"
    assert report_input.strength.classification == "strong"


def test_case_0001_html_pdf_source_shows_canonical_score() -> None:
    report_input = ReportInputV1Adapter().build(build_case_0001_source())
    html = render_html(report_input)
    assert "0.87" in html
    assert "Thân vượng" in html
    assert "51.25 / D+" not in html


def test_case_0001_docx_source_shows_canonical_score(tmp_path: Path) -> None:
    report_input = ReportInputV1Adapter().build(build_case_0001_source())
    result = ReportExportServiceV1(export_root=tmp_path).export_docx(report_input)
    validate_docx_file(Path(result.file_path))
    document = Document(result.file_path)
    paragraphs = "\n".join(paragraph.text for paragraph in document.paragraphs)
    tables = "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    text = paragraphs + "\n" + tables
    assert "0.87" in text
    assert "Thân vượng" in text


def test_legacy_html_gauge_binds_strength_payload() -> None:
    source = Path("applications/customer_portal/static/js/report/report_model.js").read_text(
        encoding="utf-8"
    )
    block = source.split("function strengthGaugeValue", 1)[1].split("function ", 1)[0]
    assert 'pick(strength, ["strength_score"])' in block
    assert "body_strength_score" not in block
    assert "than_score" not in block
