"""G1-04 Report/PDF/DOCX bind canonical Temperature Điều hậu, not pattern.dieu_hau."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from engines.report_engine.adapters.report_input_v1_adapter import ReportInputV1Adapter
from engines.report_engine.exporting.docx_exporter_v1 import validate_docx_file
from engines.report_engine.rendering.html_report_v1 import render_html
from engines.report_engine.services.report_export_service_v1 import ReportExportServiceV1
from tests.report_engine.case_0001_runtime import build_case_0001_source


def test_case_0001_report_input_uses_cold_warming() -> None:
    report_input = ReportInputV1Adapter().build(build_case_0001_source())
    assert report_input.useful_god.temperature_adjustment == "cold"
    assert report_input.useful_god.temperature_adjustment != "hot"
    assert report_input.useful_god.balancing_need == "warming"
    assert "Nguyệt lệnh Sửu" in report_input.useful_god.climate_evidence
    assert report_input.useful_god.useful_god == "Chính Quan"
    assert report_input.useful_god.useful_display == "Hỏa · Đinh · Chính Quan"
    assert report_input.useful_god.climate_preference_label == "Điều hậu ưu tiên Hỏa"
    assert report_input.useful_god.climate_display == "Hỏa · Bính · Thất Sát"
    assert report_input.pattern.primary_pattern == "Chính Ấn"
    assert report_input.pattern.follow_pattern != "cold"


def test_case_0001_html_shows_dieu_hau_not_dash() -> None:
    html = render_html(ReportInputV1Adapter().build(build_case_0001_source()))
    assert "Hàn" in html
    assert "Cần ôn ấm" in html
    assert "Sinh tháng Sửu" in html
    assert "rule cli_" not in html
    assert "Điều hậu: —" not in html
    assert ">hot<" not in html


def test_case_0001_docx_shows_climate_not_pattern_dieu_hau(tmp_path: Path) -> None:
    report_input = ReportInputV1Adapter().build(build_case_0001_source())
    result = ReportExportServiceV1(export_root=tmp_path).export_docx(report_input)
    validate_docx_file(Path(result.file_path))
    document = Document(result.file_path)
    paragraphs = "\n".join(paragraph.text for paragraph in document.paragraphs)
    tables = "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    text = paragraphs + "\n" + tables
    assert "Hàn" in text
    assert "Cần ôn ấm" in text
    assert "Hỏa · Đinh · Chính Quan" in text
    assert "Điều hậu ưu tiên Hỏa" in text
    assert "Chính Ấn" in text
