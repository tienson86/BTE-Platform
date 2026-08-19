"""G1-03 Report/PDF/DOCX bind canonical Pattern Chính Ấn."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from engines.report_engine.adapters.report_input_v1_adapter import ReportInputV1Adapter
from engines.report_engine.exporting.docx_exporter_v1 import validate_docx_file
from engines.report_engine.rendering.html_report_v1 import render_html
from engines.report_engine.services.report_export_service_v1 import ReportExportServiceV1
from tests.report_engine.case_0001_runtime import build_case_0001_source


def test_case_0001_report_input_uses_chinh_an() -> None:
    report_input = ReportInputV1Adapter().build(build_case_0001_source())
    assert report_input.pattern.primary_pattern == "Chính Ấn"
    assert report_input.pattern.primary_pattern != "Chính Quan"


def test_case_0001_html_pdf_source_shows_chinh_an() -> None:
    html = render_html(ReportInputV1Adapter().build(build_case_0001_source()))
    assert "Chính Ấn" in html
    assert "Cách chính" in html


def test_case_0001_docx_source_shows_chinh_an(tmp_path: Path) -> None:
    report_input = ReportInputV1Adapter().build(build_case_0001_source())
    result = ReportExportServiceV1(export_root=tmp_path).export_docx(report_input)
    validate_docx_file(Path(result.file_path))
    document = Document(result.file_path)
    paragraphs = "\n".join(paragraph.text for paragraph in document.paragraphs)
    tables = "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    text = paragraphs + "\n" + tables
    assert "Chính Ấn" in text
