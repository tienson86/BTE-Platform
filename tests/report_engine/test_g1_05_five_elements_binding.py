"""G1-05 Report/PDF/DOCX bind structural Five Elements counts, not Score grade."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from applications.api.services.five_elements_truth import FIVE_ELEMENTS_DISCLAIMER
from applications.api.services.orchestrator import OrchestratorService
from engines.report_engine.adapters.report_input_v1_adapter import ReportInputV1Adapter
from engines.report_engine.exporting.docx_exporter_v1 import validate_docx_file
from engines.report_engine.rendering.html_report_v1 import render_html
from engines.report_engine.rendering.report_sections_v1 import build_presented_report
from engines.report_engine.services.report_export_service_v1 import ReportExportServiceV1
from tests.report_engine.case_0001_runtime import build_case_0001_source


_FORBIDDEN_DISTRIBUTION_LABELS = (
    "Mạnh",
    "Yếu",
    "Vượng",
    "Suy",
    "Thiếu",
    "Dư",
    "nổi",
    "khuyết",
    "Dụng thần",
    "Hỷ thần",
)


def _five_elements_html(html: str) -> str:
    start = html.find("03. Phân bố Ngũ hành")
    end = html.find("04. Thân")
    assert start >= 0
    assert end > start
    return html[start:end]


def test_case_0001_report_input_uses_structural_counts() -> None:
    source = build_case_0001_source()
    report_input = ReportInputV1Adapter().build(source)
    payload = OrchestratorService().analyze(
        year=1987, month=1, day=21, hour=4, minute=30, gender="male"
    )
    counts = payload["five_elements"]["counts"]
    assert report_input.five_elements.wood == counts["wood"] == 4
    assert report_input.five_elements.fire == counts["fire"] == 5
    assert report_input.five_elements.earth == counts["earth"] == 6
    assert report_input.five_elements.metal == counts["metal"] == 3
    assert report_input.five_elements.water == counts["water"] == 1
    assert float(payload["score"]["wuxing_score"]) == 0.0
    assert report_input.five_elements.wood != payload["score"]["wuxing_score"]


def test_case_0001_html_pdf_source_shows_distribution_not_score_grade() -> None:
    report_input = ReportInputV1Adapter().build(build_case_0001_source())
    html = render_html(report_input)
    section = _five_elements_html(html)
    assert "03. Phân bố Ngũ hành" in section
    assert "Mộc" in section and ">4<" in html
    assert "Hỏa" in section
    assert "Thổ" in section
    assert "Kim" in section
    assert "Thủy" in section
    assert "Tính theo Thiên can · bản hành Địa chi · Tàng can" in section
    assert "Tổng đơn vị cấu trúc: 19" in section
    assert FIVE_ELEMENTS_DISCLAIMER in section
    table_start = section.find("<table")
    table_end = section.find("</table>")
    table = section[table_start:table_end] if table_start >= 0 and table_end > table_start else ""
    for token in _FORBIDDEN_DISTRIBUTION_LABELS:
        assert token not in table


def test_presented_section_does_not_use_score_series() -> None:
    presented = build_presented_report(
        ReportInputV1Adapter().build(build_case_0001_source())
    )
    section = next(item for item in presented.sections if item.id == "five-elements")
    rows = {row[0]: row[1] for row in (section.table.rows if section.table else [])}
    assert rows == {"Mộc": "4", "Hỏa": "5", "Thổ": "6", "Kim": "3", "Thủy": "1"}
    assert section.title == "03. Phân bố Ngũ hành"
    assert FIVE_ELEMENTS_DISCLAIMER in (section.notes or [])


def test_case_0001_docx_shows_same_distribution(tmp_path: Path) -> None:
    report_input = ReportInputV1Adapter().build(build_case_0001_source())
    result = ReportExportServiceV1(export_root=tmp_path).export_docx(report_input)
    validate_docx_file(Path(result.file_path))
    document = Document(result.file_path)
    tables = "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    paragraphs = "\n".join(paragraph.text for paragraph in document.paragraphs)
    text = paragraphs + "\n" + tables
    assert "Phân bố Ngũ hành" in text
    assert "Mộc" in text and "4" in text
    assert "Hỏa" in text and "5" in text
    assert "Thổ" in text and "6" in text
    assert "Kim" in text and "3" in text
    assert "Thủy" in text and "1" in text
    assert "Tính theo Thiên can · bản hành Địa chi · Tàng can" in text
    assert "Tổng đơn vị cấu trúc: 19" in text
    assert FIVE_ELEMENTS_DISCLAIMER in text
