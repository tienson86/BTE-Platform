"""G1-06 Report/PDF/DOCX bind rich Useful God display, keep Điều hậu separate."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from applications.api.services.orchestrator import OrchestratorService
from engines.report_engine.adapters.report_input_v1_adapter import ReportInputV1Adapter
from engines.report_engine.exporting.docx_exporter_v1 import validate_docx_file
from engines.report_engine.rendering.html_report_v1 import render_html
from engines.report_engine.services.report_export_service_v1 import ReportExportServiceV1
from tests.report_engine.case_0001_runtime import build_case_0001_source

CANONICAL_DISPLAY = "Hỏa · Đinh · Chính Quan"
FAVORABLE_DISPLAY = "Hỏa · Đinh · Chính Quan / Thủy · Nhâm · Thực Thần"
UNFAVORABLE_DISPLAY = "Kim · Canh · Tỷ Kiên / Kim · Tân · Kiếp Tài"
CLIMATE_PREFERENCE = "Điều hậu ưu tiên Hỏa"


def test_case_0001_report_uses_rich_useful_god_not_dieu_hau() -> None:
    report_input = ReportInputV1Adapter().build(build_case_0001_source())
    useful = report_input.useful_god
    assert useful.useful_god == "Chính Quan"
    assert useful.useful_display == CANONICAL_DISPLAY
    assert useful.useful_ten_god == "Chính Quan"
    assert useful.useful_stem == "Đinh"
    assert useful.useful_element == "Hỏa"
    assert useful.winning_rule_id == "str_003"
    assert useful.winning_rule_group == "strength"
    assert useful.climate_rule_id == "sea_001"
    assert useful.climate_preference_label == CLIMATE_PREFERENCE
    assert useful.favorable_display == FAVORABLE_DISPLAY
    assert useful.unfavorable_display == UNFAVORABLE_DISPLAY
    assert useful.temperature_adjustment == "cold"
    assert useful.balancing_need == "warming"
    assert useful.useful_god != useful.balancing_need
    assert useful.useful_display != "Cần ôn ấm"
    assert useful.useful_display != useful.climate_display
    assert report_input.pattern.primary_pattern == "Chính Ấn"


def test_case_0001_api_and_report_same_useful_god() -> None:
    payload = OrchestratorService().analyze(
        year=1987, month=1, day=21, hour=4, minute=30, gender="male"
    )
    report_input = ReportInputV1Adapter().build(build_case_0001_source())
    useful = payload["useful_god"]
    report = report_input.useful_god
    assert useful["useful_display"] == report.useful_display == CANONICAL_DISPLAY
    assert useful["winning_rule_id"] == report.winning_rule_id == "str_003"
    assert useful["climate_rule_id"] == report.climate_rule_id == "sea_001"
    assert useful["favorable_display"] == report.favorable_display
    assert useful["unfavorable_display"] == report.unfavorable_display
    assert payload["temperature"]["climate_state"] == report.temperature_adjustment == "cold"
    assert payload["temperature"]["balancing_need"] == report.balancing_need == "warming"


def test_case_0001_html_shows_three_layer_useful_god(tmp_path: Path) -> None:
    report_input = ReportInputV1Adapter().build(build_case_0001_source())
    html = render_html(report_input)
    assert CANONICAL_DISPLAY in html
    assert FAVORABLE_DISPLAY in html
    assert UNFAVORABLE_DISPLAY in html
    assert CLIMATE_PREFERENCE in html
    assert "Hàn" in html
    assert "Cần ôn ấm" in html
    result = ReportExportServiceV1(export_root=tmp_path).export_docx(report_input)
    validate_docx_file(Path(result.file_path))
    document = Document(result.file_path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    tables = "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    blob = text + "\n" + tables
    assert CANONICAL_DISPLAY in blob
    assert "Hàn" in blob
    assert "Cần ôn ấm" in blob
    assert "Chính Ấn" in blob
