"""CASE-0001 PDF/DOCX integration export tests."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from engines.report_engine.adapters.report_input_v1_adapter import ReportInputV1Adapter
from engines.report_engine.exporting.docx_exporter_v1 import validate_docx_file
from engines.report_engine.exporting.pdf_exporter_v1 import validate_pdf_file
from engines.report_engine.rendering.html_report_v1 import render_html
from engines.report_engine.services.report_export_service_v1 import ReportExportServiceV1
from tests.report_engine.case_0001_runtime import CASE_0001_CANONICAL, build_case_0001_source

EXPORT_DIR = Path("knowledge/report_v1_validation/exports")
PDF_NAME = "BTE_CASE-0001_Nguyen_Tien_Son_Report_V1_0.pdf"
DOCX_NAME = "BTE_CASE-0001_Nguyen_Tien_Son_Report_V1_0.docx"

VIETNAMESE_SMOKE_STRINGS = (
    "Nguyễn Tiến Sơn",
    "Bát Tự",
    "Thân vượng nhược",
    "Dụng thần",
    "Hỷ thần",
    "Kỵ thần",
)


@pytest.fixture(scope="module")
def case_0001_report_input():
    """Build CASE-0001 ReportInputV1 once per module."""
    return ReportInputV1Adapter().build(build_case_0001_source())


def test_case_0001_integration_exports(case_0001_report_input, tmp_path_factory) -> None:
    """CASE-0001 fixture → adapter → export service → PDF + DOCX."""
    export_root = tmp_path_factory.mktemp("case_0001_exports")
    service = ReportExportServiceV1(export_root=export_root)
    pdf_result = service.export_pdf(case_0001_report_input)
    docx_result = service.export_docx(case_0001_report_input)
    validate_pdf_file(Path(pdf_result.file_path))
    validate_docx_file(Path(docx_result.file_path))
    assert pdf_result.case_id == "CASE-0001"
    assert docx_result.case_id == "CASE-0001"
    assert pdf_result.size_bytes > 0
    assert docx_result.size_bytes > 0


def test_case_0001_vietnamese_docx_smoke(case_0001_report_input, tmp_path: Path) -> None:
    """DOCX contains required Vietnamese smoke strings."""
    service = ReportExportServiceV1(export_root=tmp_path)
    result = service.export_docx(case_0001_report_input)
    document = Document(result.file_path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Nguyễn Tiến Sơn" in text
    assert "BÁO CÁO LUẬN GIẢI BÁT TỰ" in text
    for token in ("Thân vượng nhược", "Dụng thần"):
        assert token in text


def test_case_0001_html_unicode_before_pdf(case_0001_report_input) -> None:
    """HTML source used for PDF contains Vietnamese text."""
    html = render_html(case_0001_report_input)
    for token in ("Nguyễn Tiến Sơn", "Bát Tự", "Canh", "Bính Dần"):
        assert token in html


def test_case_0001_generate_artifacts(case_0001_report_input) -> None:
    """Generate CASE-0001 PDF/DOCX artifacts under validation exports."""
    service = ReportExportServiceV1(export_root=EXPORT_DIR)
    pdf_result = service.export_pdf(
        case_0001_report_input,
        EXPORT_DIR / PDF_NAME,
    )
    docx_result = service.export_docx(
        case_0001_report_input,
        EXPORT_DIR / DOCX_NAME,
    )
    pdf_path = Path(pdf_result.file_path)
    docx_path = Path(docx_result.file_path)
    assert pdf_path.name == PDF_NAME
    assert docx_path.name == DOCX_NAME
    assert pdf_path.stat().st_size > 10_000
    assert docx_path.stat().st_size > 5_000
    validate_pdf_file(pdf_path)
    validate_docx_file(docx_path)
    document = Document(str(docx_path))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert CASE_0001_CANONICAL["profile"].full_name in text
