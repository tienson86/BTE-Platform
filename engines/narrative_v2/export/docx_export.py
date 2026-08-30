"""Shadow DOCX export. Presentation blocks only. No Report Builder."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.shared import Pt

from engines.narrative_v2.export.export_context import ExportBlock, ExportContext


@dataclass(frozen=True, slots=True)
class DocxExport:
    """Shadow DOCX plus extracted paragraph texts."""

    version: str
    status: str
    blocks: tuple[ExportBlock, ...]
    paragraphs: tuple[str, ...]
    path: str | None


def export_docx(context: ExportContext, output_path: Path | None = None) -> DocxExport:
    """Write one paragraph per Presentation block. No extra wording."""
    if output_path is None:
        raise TypeError("output_path is required for DOCX export")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    core = document.core_properties
    core.title = context.version
    core.subject = "narrative-v2-shadow"
    for block in context.blocks:
        document.add_paragraph(block.text)
    document.save(str(output_path))
    paragraphs = extract_docx_paragraphs(output_path)
    return DocxExport(
        version=context.version,
        status=context.status,
        blocks=context.blocks,
        paragraphs=paragraphs,
        path=str(output_path.resolve()),
    )


def extract_docx_paragraphs(path: Path) -> tuple[str, ...]:
    """Read customer paragraphs back from the saved DOCX."""
    document = Document(str(path))
    return tuple(paragraph.text for paragraph in document.paragraphs if paragraph.text)
