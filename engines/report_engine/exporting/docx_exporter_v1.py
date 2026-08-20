"""DOCX Export V1 — editable Word report from ReportInputV1."""

from __future__ import annotations

import logging
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Cm, Pt

from engines.report_engine.contracts.report_export_result_v1 import (
    EXPORT_FORMAT_DOCX,
    MEDIA_TYPE_DOCX,
    ReportExportResultV1,
)
from engines.report_engine.contracts.report_input_v1 import ReportInputV1
from engines.report_engine.rendering.report_sections_v1 import (
    PresentedSection,
    build_presented_report,
)

logger = logging.getLogger(__name__)

DOCX_MIN_BYTES = 2048


class DocxExporterV1:
    """Export ReportInputV1 to a real OpenXML DOCX document."""

    def export(
        self,
        report_input: ReportInputV1,
        output_path: Path,
    ) -> ReportExportResultV1:
        """Build and save DOCX to output_path."""
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document = Document()
        self._configure_page(document)
        self._apply_styles(document)
        self._stamp_properties(document, report_input)
        self._render(document, report_input)
        document.save(str(output_path))
        validate_docx_file(output_path)
        logger.info(
            "report_export_docx case_id=%s path=%s size=%s",
            report_input.metadata.case_id,
            output_path,
            output_path.stat().st_size,
        )
        return ReportExportResultV1(
            format=EXPORT_FORMAT_DOCX,
            file_path=str(output_path.resolve()),
            file_name=output_path.name,
            media_type=MEDIA_TYPE_DOCX,
            size_bytes=output_path.stat().st_size,
            report_version=report_input.metadata.report_version,
            case_id=report_input.metadata.case_id,
            generated_at=report_input.metadata.generated_at,
            page_count=None,
        )

    def _configure_page(self, document: Document) -> None:
        section = document.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    def _apply_styles(self, document: Document) -> None:
        normal = document.styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(11)

    def _stamp_properties(self, document: Document, report_input: ReportInputV1) -> None:
        """Embed identity in document properties — not a second analytical model."""
        props = document.core_properties
        props.title = "Báo cáo luận giải Bát Tự"
        props.subject = report_input.profile.full_name or "BTE V1.0"
        props.identifier = report_input.metadata.case_id
        props.comments = (
            f"BTE V1.0 · Report V{report_input.metadata.report_version} · "
            f"analysis_id={report_input.metadata.case_id}"
        )

    def _render(self, document: Document, report_input: ReportInputV1) -> None:
        presented = build_presented_report(report_input)
        document.add_heading("BÁO CÁO LUẬN GIẢI BÁT TỰ", level=0)
        subtitle_name = report_input.profile.full_name or "—"
        document.add_paragraph(subtitle_name)
        document.add_paragraph(presented.subtitle)
        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        for section in presented.sections:
            self._render_section(document, section)
        document.add_paragraph()
        document.add_paragraph(presented.footer, style="Caption")

    def _render_section(self, document: Document, section: PresentedSection) -> None:
        document.add_heading(section.title, level=1)
        if section.meta_rows:
            self._key_value_table(document, section.meta_rows)
        if section.table is not None:
            self._grid_table(document, section.table.headers, section.table.rows)
        for paragraph in section.paragraphs:
            document.add_paragraph(paragraph)
        for item in section.list_items:
            document.add_paragraph(item, style="List Bullet")
        for note in section.notes:
            document.add_paragraph(note)
        has_content = bool(
            section.meta_rows
            or section.table
            or section.paragraphs
            or section.list_items
            or section.pillars
        )
        if section.fallback and not has_content:
            document.add_paragraph(section.fallback)

    def _key_value_table(
        self,
        document: Document,
        rows: list[tuple[str, str]],
    ) -> None:
        table = document.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for label, value in rows:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = value

    def _grid_table(
        self,
        document: Document,
        headers: list[str],
        rows: list[list[str]],
    ) -> None:
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        for index, label in enumerate(headers):
            header_cells[index].text = label
        for row in rows:
            cells = table.add_row().cells
            for index, cell in enumerate(row):
                cells[index].text = cell


def export_docx(report_input: ReportInputV1, output_path: Path) -> ReportExportResultV1:
    """Module-level DOCX export helper."""
    return DocxExporterV1().export(report_input, output_path)


def validate_docx_file(output_path: Path) -> None:
    """Validate DOCX is a non-empty OpenXML zip archive."""
    if not output_path.is_file():
        raise FileNotFoundError(f"DOCX not created: {output_path}")
    if not zipfile.is_zipfile(output_path):
        raise ValueError(f"Invalid DOCX zip archive: {output_path}")
    size = output_path.stat().st_size
    if size < DOCX_MIN_BYTES:
        raise ValueError(f"DOCX too small ({size} bytes): {output_path}")
    with zipfile.ZipFile(output_path, "r") as archive:
        if "[Content_Types].xml" not in archive.namelist():
            raise ValueError(f"Invalid DOCX content types: {output_path}")
