"""Date Selection DOCX exporter. Reuses PACK 05 python-docx pipeline."""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from engines.date_selection_report.exceptions import DateSelectionReportExportError
from engines.date_selection_report.exporting.filename import build_docx_filename
from engines.date_selection_report.exporting.html_projection import (
    PDF_AUTHOR,
    PDF_DOCUMENT_TITLE,
    PDF_SUBJECT,
)
from engines.date_selection_report.rendering.labels import FORBIDDEN_PUBLIC_TERMS
from engines.date_selection_report.rendering.nodes import (
    DateSelectionRenderTree,
    RecommendationNode,
)
from engines.report_engine.contracts.report_export_result_v1 import (
    EXPORT_FORMAT_DOCX,
    MEDIA_TYPE_DOCX,
    ReportExportResultV1,
)
from engines.report_engine.exporting.docx_exporter_v1 import DocxExporterV1, validate_docx_file

logger = logging.getLogger(__name__)

_ACCENT = RGBColor(0x2C, 0x4A, 0x6E)
_MUTED = RGBColor(0x55, 0x55, 0x55)
_NAMED_STYLES: tuple[tuple[str, int, bool, RGBColor], ...] = (
    ("ReportTitle", 22, True, _ACCENT),
    ("SectionTitle", 12, True, _ACCENT),
    ("RecommendationTitle", 16, True, _ACCENT),
    ("Result", 13, True, _ACCENT),
    ("Label", 10, False, _MUTED),
    ("Value", 11, True, RGBColor(0x1A, 0x1A, 0x1A)),
    ("Caption", 10, False, _MUTED),
    ("Footer", 9, False, _MUTED),
)


class DateSelectionDocxExporter(DocxExporterV1):
    """Render a frozen DateSelectionRenderTree through PACK 05 DOCX primitives."""

    def export(
        self,
        tree: DateSelectionRenderTree,
        output_path: Path,
    ) -> ReportExportResultV1:
        """Write an editable OpenXML document from the render tree."""
        _require_tree(tree)
        target = _resolve_docx_path(tree, output_path)
        target.parent.mkdir(parents=True, exist_ok=True)
        document = Document()
        self._configure_page(document)
        self._apply_styles(document)
        _set_east_asia(document.styles["Normal"], "Arial")
        self._add_named_styles(document)
        self._render_tree(document, tree)
        self._apply_core_properties(document, tree)
        assert_docx_text_clean(extract_document_text(document))
        document.save(str(target))
        validate_docx_file(target)
        logger.info(
            "date_selection_docx report_id=%s path=%s size=%s",
            tree.header.report_id,
            target,
            target.stat().st_size,
        )
        return ReportExportResultV1(
            format=EXPORT_FORMAT_DOCX,
            file_path=str(target.resolve()),
            file_name=target.name,
            media_type=MEDIA_TYPE_DOCX,
            size_bytes=target.stat().st_size,
            report_version=tree.footer.report_version,
            case_id=tree.header.report_id,
            generated_at=tree.header.generated_at,
            page_count=None,
        )

    def _add_named_styles(self, document: Document) -> None:
        existing = {item.name for item in document.styles}
        for name, size, bold, color in _NAMED_STYLES:
            if name in existing:
                style = document.styles[name]
            else:
                style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
                style.base_style = document.styles["Normal"]
            style.font.name = "Arial"
            style.font.size = Pt(size)
            style.font.bold = bold
            try:
                style.font.color.rgb = color
            except (ValueError, AttributeError):
                pass
            _set_east_asia(style, "Arial")

    def _render_tree(self, document: Document, tree: DateSelectionRenderTree) -> None:
        self._cover(document, tree)
        self._executive(document, tree)
        self._person(document, tree)
        self._search(document, tree)
        self._recommendations(document, tree)
        self._guidance(document, tree)
        self._footer_block(document, tree)

    def _cover(self, document: Document, tree: DateSelectionRenderTree) -> None:
        document.add_paragraph(tree.header.title, style="ReportTitle")
        document.add_paragraph(tree.header.subtitle, style="Caption")
        document.add_paragraph(_format_generated_at(tree.header.generated_at), style="Caption")

    def _executive(self, document: Document, tree: DateSelectionRenderTree) -> None:
        period = tree.search_period
        rows = [
            ("Khách hàng", _person_value(tree, "full_name")),
            ("Nhóm Trạch", _person_value(tree, "trach_group")),
            (period.month_label, period.month_display),
            (period.recommendation_count_label, period.recommendation_count),
        ]
        document.add_paragraph("Tóm tắt", style="SectionTitle")
        self._key_value_table(document, rows)

    def _person(self, document: Document, tree: DateSelectionRenderTree) -> None:
        document.add_paragraph(tree.person.title, style="SectionTitle")
        self._paired_table(document, [(row.label, row.value) for row in tree.person.rows])

    def _search(self, document: Document, tree: DateSelectionRenderTree) -> None:
        period = tree.search_period
        document.add_paragraph(period.title, style="SectionTitle")
        rows = [
            (period.month_label, period.month_display),
            (period.recommendation_count_label, period.recommendation_count),
        ]
        self._key_value_table(document, rows)
        if period.explanation:
            document.add_paragraph(period.explanation, style="Caption")

    def _recommendations(self, document: Document, tree: DateSelectionRenderTree) -> None:
        document.add_paragraph(tree.recommendations_title, style="SectionTitle")
        if tree.empty_state is not None:
            document.add_paragraph(tree.empty_state.message, style="Caption")
            return
        for node in tree.recommendations:
            self._recommendation(document, node)

    def _recommendation(self, document: Document, node: RecommendationNode) -> None:
        header = node.date_header
        rank = document.add_paragraph(f"{node.rank:02d}", style="Caption")
        date = document.add_paragraph(header.solar_date, style="RecommendationTitle")
        lunar = document.add_paragraph(header.lunar_display, style="Caption")
        result = document.add_paragraph(header.day_result, style="Result")
        for paragraph in (rank, date, lunar, result):
            paragraph.paragraph_format.keep_with_next = True
        self._key_value_table(
            document,
            [(row.label, row.value) for row in node.day_information.rows],
        )
        hours_title = document.add_paragraph(
            node.compatible_hours.title,
            style="SectionTitle",
        )
        hours_title.paragraph_format.keep_with_next = True
        for hour in node.compatible_hours.rows:
            document.add_paragraph(hour.display, style="List Bullet")
        times_title = document.add_paragraph(
            node.positive_times.title,
            style="SectionTitle",
        )
        times_title.paragraph_format.keep_with_next = True
        for group in node.positive_times.groups:
            heading = document.add_paragraph(group.label, style="Value")
            heading.paragraph_format.keep_with_next = True
            for item in group.items:
                document.add_paragraph(
                    f"{item.branch_display} · {item.time_range}",
                    style="List Bullet",
                )

    def _guidance(self, document: Document, tree: DateSelectionRenderTree) -> None:
        guidance_title = document.add_paragraph(tree.guidance.title, style="SectionTitle")
        guidance_title.paragraph_format.keep_with_next = True
        for item in tree.guidance.items:
            label = document.add_paragraph(item.label, style="Value")
            label.paragraph_format.keep_with_next = True
            document.add_paragraph(item.text)

    def _footer_block(self, document: Document, tree: DateSelectionRenderTree) -> None:
        footer = tree.footer
        text = f"{footer.generator} · {footer.product} · {footer.report_version}"
        document.add_paragraph(text, style="Footer")

    def _key_value_table(
        self,
        document: Document,
        rows: list[tuple[str, str]],
    ) -> None:
        table = document.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for label, value in rows:
            cells = table.add_row().cells
            _set_cell(cells[0], label, "Label")
            _set_cell(cells[1], value, "Value")

    def _paired_table(self, document: Document, rows: list[tuple[str, str]]) -> None:
        table = document.add_table(rows=0, cols=4)
        table.style = "Table Grid"
        pairs = list(rows)
        for index in range(0, len(pairs), 2):
            cells = table.add_row().cells
            _set_cell(cells[0], pairs[index][0], "Label")
            _set_cell(cells[1], pairs[index][1], "Value")
            if index + 1 < len(pairs):
                _set_cell(cells[2], pairs[index + 1][0], "Label")
                _set_cell(cells[3], pairs[index + 1][1], "Value")

    def _apply_core_properties(self, document: Document, tree: DateSelectionRenderTree) -> None:
        props = document.core_properties
        props.title = PDF_DOCUMENT_TITLE
        props.author = PDF_AUTHOR
        props.subject = PDF_SUBJECT
        props.identifier = tree.header.report_id


def extract_document_text(document: Document) -> str:
    """Collect paragraph and table text for validation and tests."""
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def extract_docx_text(path: Path) -> str:
    """Reopen a DOCX and return concatenated text."""
    return extract_document_text(Document(str(path)))


def export_docx(tree: DateSelectionRenderTree, output_path: Path) -> ReportExportResultV1:
    """Module-level Date Selection DOCX export helper."""
    return DateSelectionDocxExporter().export(tree, output_path)


def _set_cell(cell: object, text: str, style_name: str) -> None:
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph.style = style_name


def _set_east_asia(style: object, name: str) -> None:
    element = getattr(style, "element", None)
    if element is None:
        return
    r_pr = element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), name)


def _person_value(tree: DateSelectionRenderTree, key: str) -> str:
    for row in tree.person.rows:
        if row.key == key:
            return row.value
    return ""


def _format_generated_at(value: str) -> str:
    date_part = value.split("T", 1)[0]
    pieces = date_part.split("-")
    if len(pieces) == 3:
        year, month, day = pieces
        return f"{day}/{month}/{year}"
    return value


def _require_tree(tree: DateSelectionRenderTree) -> None:
    if tree is None:
        raise DateSelectionReportExportError("RenderTree is required")
    if not tree.header.title:
        raise DateSelectionReportExportError("missing header")
    if not tree.person.rows:
        raise DateSelectionReportExportError("missing person")


def _resolve_docx_path(tree: DateSelectionRenderTree, output_path: Path) -> Path:
    if output_path.suffix.lower() == ".docx":
        return output_path
    return output_path / build_docx_filename(tree)


def assert_docx_text_clean(text: str) -> None:
    """Abort when placeholders or forbidden hour terminology leak into DOCX."""
    if "{{" in text or "}}" in text:
        raise DateSelectionReportExportError("unresolved placeholders remain")
    for term in FORBIDDEN_PUBLIC_TERMS:
        if term in text:
            raise DateSelectionReportExportError(f"forbidden public term: {term}")
