"""
BTE Platform
DOCX Renderer
"""

from __future__ import annotations

from io import BytesIO

from docx import Document

from .base_renderer import BaseRenderer

from ..models import Chapter


class DocxRenderer(BaseRenderer):

    name = "docx"

    extension = ".docx"

    mime_type = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    def render(
        self,
        chapters: list[Chapter],
    ) -> bytes:

        document = Document()

        for chapter in chapters:

            document.add_heading(
                chapter.title,
                level=1,
            )

            for paragraph in chapter.paragraphs:

                document.add_heading(
                    paragraph.title,
                    level=2,
                )

                for sentence in paragraph.sentences:

                    document.add_paragraph(
                        sentence.text
                    )

        buffer = BytesIO()

        document.save(buffer)

        return buffer.getvalue()
